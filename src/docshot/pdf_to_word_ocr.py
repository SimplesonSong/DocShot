"""OCR-based PDF to Word conversion workflow."""

from __future__ import annotations

from collections.abc import Callable
from gc import collect
from pathlib import Path
from tempfile import TemporaryDirectory

from app.ocr_engine import OcrEngine


class PdfToWordOcrError(Exception):
    """Raised when OCR-mode PDF to Word conversion fails."""


_OCR_RENDER_ZOOM = 1.5


def pdf_to_word_ocr(
    pdf_path: str,
    output_docx_path: str,
    status_callback: Callable[[str], None] | None = None,
    progress_callback: Callable[[int, int], None] | None = None,
) -> str:
    """Convert PDF to Word via OCR and return output DOCX path."""
    source = Path(pdf_path)
    target = Path(output_docx_path)

    if not source.exists():
        raise FileNotFoundError(f"PDF file does not exist: {source}")

    if source.suffix.lower() != ".pdf":
        raise ValueError(
            f"Unsupported file extension: {source.suffix or '<none>'}. Expected extension: .pdf"
        )

    target.parent.mkdir(parents=True, exist_ok=True)

    doc = None
    try:
        import fitz
        from docx import Document

        if status_callback is not None:
            status_callback("正在初始化 OCR")

        ocr = OcrEngine()
        docx = Document()
        doc = fitz.open(str(source))
        total_pages = len(doc)

        with TemporaryDirectory(prefix="docshot_ocr_") as temp_dir:
            temp_root = Path(temp_dir)

            for idx, page in enumerate(doc, start=1):
                temp_image = temp_root / f"page_{idx:04d}.png"
                if status_callback is not None:
                    status_callback(f"正在渲染第 {idx} 页")

                try:
                    pix = page.get_pixmap(
                        matrix=fitz.Matrix(_OCR_RENDER_ZOOM, _OCR_RENDER_ZOOM),
                        alpha=False,
                        annots=False,
                    )
                    pix.save(str(temp_image))
                    del pix

                    if status_callback is not None:
                        status_callback(f"正在识别第 {idx} 页")

                    texts = ocr.recognize_image(str(temp_image))
                finally:
                    if temp_image.exists():
                        temp_image.unlink(missing_ok=True)
                    collect()

                if status_callback is not None:
                    status_callback("正在写入 Word")

                if texts:
                    for line in texts:
                        docx.add_paragraph(line)
                else:
                    docx.add_paragraph("")

                if idx < total_pages:
                    docx.add_page_break()

                if progress_callback is not None:
                    progress_callback(idx, total_pages)

        docx.save(str(target))

        if progress_callback is not None:
            progress_callback(total_pages, total_pages)
        if status_callback is not None:
            status_callback("转换完成")

        return str(target)
    except Exception as exc:  # pragma: no cover - depends on local files/environment
        raise PdfToWordOcrError(
            f"Failed to convert PDF to Word (OCR mode). source='{source}', target='{target}', "
            f"error='{exc}'"
        ) from exc
    finally:
        if doc is not None:
            doc.close()
